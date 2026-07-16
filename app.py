import streamlit as st
import time
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq

st.set_page_config(page_title="HR Policy Assistant (RAG)", page_icon="📄", layout="wide")

# ==================================================================
# Sidebar: API key + file upload
# ==================================================================
st.sidebar.title("Setup")

# Check if a key was set via Streamlit Secrets (safe, not stored in code)
default_key = st.secrets.get("GROQ_API_KEY", "") if hasattr(st, "secrets") else ""

if default_key:
    groq_key = default_key
    st.sidebar.success("Groq API key loaded from Secrets ✓")
else:
    groq_key = st.sidebar.text_input("Groq API Key", type="password", help="Get a free key at console.groq.com")

uploaded_files = st.sidebar.file_uploader(
    "Upload document(s)", type=["txt", "pdf", "docx"], accept_multiple_files=True
)

process_btn = st.sidebar.button("Process Document(s)", use_container_width=True)

st.sidebar.markdown("---")
st.sidebar.caption(
    "This app builds a RAG (Retrieval Augmented Generation) pipeline: "
    "your document is chunked, embedded, stored in a FAISS vector database, "
    "then retrieved and passed to Llama 3 (via Groq) to answer your questions."
)

# ==================================================================
# Cached model loaders (so they only load once, not on every interaction)
# ==================================================================
@st.cache_resource
def get_embedding_model():
    return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")


def get_llm(api_key):
    return ChatGroq(model="llama-3.1-8b-instant", temperature=0.2, groq_api_key=api_key)


# ==================================================================
# Document extraction
# ==================================================================
def extract_text(uploaded_file):
    ext = uploaded_file.name.lower().split(".")[-1]
    if ext == "txt":
        return uploaded_file.read().decode("utf-8")
    elif ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text
    elif ext == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return ""


# ==================================================================
# Session state init
# ==================================================================
if "vector_db" not in st.session_state:
    st.session_state.vector_db = None
if "messages" not in st.session_state:
    st.session_state.messages = []
if "doc_summary" not in st.session_state:
    st.session_state.doc_summary = None

# ==================================================================
# Process documents on button click
# ==================================================================
if process_btn:
    if not groq_key:
        st.sidebar.error("Please enter your Groq API key.")
    elif not uploaded_files:
        st.sidebar.error("Please upload at least one document.")
    else:
        with st.spinner("Processing document(s)..."):
            full_text = ""
            file_info = []
            for f in uploaded_files:
                text = extract_text(f)
                if text.strip():
                    full_text += text + "\n"
                    file_info.append(f"{f.name} ({len(text)} chars)")
                else:
                    st.sidebar.warning(f"No text extracted from {f.name}")

            if full_text.strip():
                splitter = RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=50)
                chunks = splitter.create_documents([full_text])

                embedding_model = get_embedding_model()
                vector_db = FAISS.from_documents(chunks, embedding_model)

                st.session_state.vector_db = vector_db
                st.session_state.messages = []  # reset chat on new document
                st.session_state.doc_summary = {
                    "files": file_info,
                    "num_chunks": len(chunks),
                    "sample_chunks": [c.page_content for c in chunks[:3]],
                }
                st.sidebar.success(f"Processed {len(uploaded_files)} file(s) into {len(chunks)} chunks.")
            else:
                st.sidebar.error("No usable text found in uploaded file(s).")

# ==================================================================
# Main area
# ==================================================================
st.title("📄 HR Policy Assistant — RAG Chatbot")
st.caption("Upload an HR policy document, then ask questions about it in plain English.")

if st.session_state.doc_summary:
    with st.expander("Document processing details (chunks, files)", expanded=False):
        st.write("**Files loaded:**")
        for f in st.session_state.doc_summary["files"]:
            st.write(f"- {f}")
        st.write(f"**Total chunks created:** {st.session_state.doc_summary['num_chunks']}")
        st.write("**Sample chunks:**")
        for i, c in enumerate(st.session_state.doc_summary["sample_chunks"]):
            st.code(c, language=None)

# Display chat history
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])
        if msg["role"] == "assistant" and "details" in msg:
            with st.expander("Retrieval + eval details"):
                st.write(f"**Retrieval latency:** {msg['details']['retrieval_time']:.3f} sec")
                st.write(f"**Generation latency:** {msg['details']['gen_time']:.2f} sec")
                st.write(f"**Context overlap (grounding check):** {msg['details']['overlap']:.1%}")
                st.write("**Retrieved chunks:**")
                for i, (chunk, score) in enumerate(msg["details"]["chunks"]):
                    st.write(f"Chunk {i+1} (distance: {score:.4f})")
                    st.code(chunk, language=None)

# Chat input
question = st.chat_input("Ask a question about the uploaded document...")

if question:
    if st.session_state.vector_db is None:
        st.error("Please upload and process a document first (use the sidebar).")
    else:
        st.session_state.messages.append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.write(question)

        with st.chat_message("assistant"):
            with st.spinner("Retrieving and generating answer..."):
                t0 = time.time()
                results = st.session_state.vector_db.similarity_search_with_score(question, k=3)
                retrieval_time = time.time() - t0

                context = "\n\n".join([doc.page_content for doc, _ in results])
                prompt = f"""Answer the question using ONLY the context below. If the answer isn't in the context, say so clearly.

Context:
{context}

Question: {question}
Answer:"""

                llm = get_llm(groq_key)
                t0 = time.time()
                response = llm.invoke(prompt)
                gen_time = time.time() - t0
                answer = response.content

                answer_words = set(answer.lower().split())
                context_words = set(context.lower().split())
                overlap = len(answer_words & context_words) / max(len(answer_words), 1)

                st.write(answer)

                details = {
                    "retrieval_time": retrieval_time,
                    "gen_time": gen_time,
                    "overlap": overlap,
                    "chunks": [(doc.page_content, score) for doc, score in results],
                }

                with st.expander("Retrieval + eval details"):
                    st.write(f"**Retrieval latency:** {retrieval_time:.3f} sec")
                    st.write(f"**Generation latency:** {gen_time:.2f} sec")
                    st.write(f"**Context overlap (grounding check):** {overlap:.1%}")
                    st.write("**Retrieved chunks:**")
                    for i, (chunk, score) in enumerate(details["chunks"]):
                        st.write(f"Chunk {i+1} (distance: {score:.4f})")
                        st.code(chunk, language=None)

        st.session_state.messages.append({"role": "assistant", "content": answer, "details": details})
